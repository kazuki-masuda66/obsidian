import sys
import os

def extract_text_from_docx(docx_path):
    try:
        from docx import Document
        
        result = []
        result.append(f"=== DOCX File: {os.path.basename(docx_path)} ===")
        result.append("")
        
        doc = Document(docx_path)
        
        # 段落からテキストを抽出
        for para_idx, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                result.append(f"Paragraph {para_idx}: {text}")
        
        result.append("")
        
        # テーブルからテキストを抽出
        for table_idx, table in enumerate(doc.tables, start=1):
            result.append(f"--- Table {table_idx} ---")
            for row_idx, row in enumerate(table.rows, start=1):
                row_data = []
                for cell in row.cells:
                    if cell.text:
                        row_data.append(cell.text.strip())
                if row_data:
                    result.append(f"Row {row_idx}: {' | '.join(row_data)}")
            result.append("")
        
        return "\n".join(result)
    except ImportError:
        return "Error: python-docx library is not installed. Please install it using: pip install python-docx"
    except Exception as e:
        import traceback
        return f"Error extracting text from DOCX: {e}\n{traceback.format_exc()}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = extract_text_from_docx(file_path)
    
    # Ensure output is encoded in UTF-8 to handle special characters
    sys.stdout.buffer.write(result.encode('utf-8', errors='replace'))


